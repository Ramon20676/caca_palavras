import random
import unicodedata
import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt

def normalizar_palavra(palavra):
    # Remove acentos e espaços de uma palavra
    palavra = unicodedata.normalize('NFKD', palavra).encode('ASCII', 'ignore').decode('ASCII')
    return palavra.replace(" ", "")

def verificar_espaco(matriz, palavra, linha, coluna, direcao):
    # Verifica se há espaço disponível para inserir a palavra
    tamanho = len(matriz)
    if direcao == "horizontal":
        if coluna + len(palavra) > tamanho:
            return False
        return all(matriz[linha][coluna + i] in (' ', palavra[i].upper()) for i in range(len(palavra)))
    elif direcao == "vertical":
        if linha + len(palavra) > tamanho:
            return False
        return all(matriz[linha + i][coluna] in (' ', palavra[i].upper()) for i in range(len(palavra)))

def inserir_palavra(matriz, palavra, linha, coluna, direcao):
    # Insere a palavra na matriz na direção especificada
    if direcao == "horizontal":
        for i, letra in enumerate(palavra):
            matriz[linha][coluna + i] = letra.upper()
    elif direcao == "vertical":
        for i, letra in enumerate(palavra):
            matriz[linha + i][coluna] = letra.upper()

def gerar_caca_palavras(palavras):
    # Configuração do tamanho da matriz
    maior_palavra = max(len(palavra) for palavra in palavras)
    tamanho = max(10, maior_palavra + 2, len(palavras) + 2)  # Garante espaço suficiente

    # Criar matriz vazia preenchida com letras aleatórias
    matriz = [[' ' for _ in range(tamanho)] for _ in range(tamanho)]

    # Inserir palavras na matriz de forma variada
    coordenadas_palavras = []  # Para destacar as palavras
    horizontal = True  # Alternar entre horizontal e vertical
    for palavra in palavras:
        inserido = False
        for _ in range(100):  # Tentativas para encontrar posição válida
            linha = random.randint(0, tamanho - 1)
            coluna = random.randint(0, tamanho - 1)
            direcao = "horizontal" if horizontal else "vertical"
            if verificar_espaco(matriz, palavra, linha, coluna, direcao):
                inserir_palavra(matriz, palavra, linha, coluna, direcao)
                coordenadas_palavras.append((palavra, linha, coluna, direcao))
                inserido = True
                break
        horizontal = not horizontal  # Alternar direção

    # Preencher espaços vazios com letras aleatórias
    for i in range(tamanho):
        for j in range(tamanho):
            if matriz[i][j] == ' ':
                matriz[i][j] = random.choice('ABCDEFGHIJKLMNOPQRSTUVWXYZ')

    return matriz, coordenadas_palavras

def salvar_docx(matriz, palavras):
    doc = Document()
    doc.add_heading("Caça-Palavras", level=1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Consolas'
    font.size = Pt(12)

    doc.add_paragraph("Este é o caça-palavras gerado com as palavras fornecidas.")

    doc.add_heading("Versão Original:", level=2)
    for linha in matriz:
        doc.add_paragraph(" ".join(linha))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Interface do Streamlit
st.title("Gerador de Caça-Palavras")

st.write("Adicione palavras para criar seu caça-palavras. Quando terminar, clique em 'Gerar'.")

palavras_input = st.text_area("Digite as palavras (uma por linha):")

gerar = st.button("Gerar Caça-Palavras")

if gerar:
    palavras = [normalizar_palavra(p.strip()) for p in palavras_input.split("\n") if p.strip()]
    if palavras:
        matriz, coordenadas = gerar_caca_palavras(palavras)

        st.subheader("Caça-Palavras Gerado:")
        st.code("\\n".join([" ".join(linha) for linha in matriz]), language="plaintext")

        docx_file = salvar_docx(matriz, palavras)
        st.download_button(
            label="Baixar Caça-Palavras (.docx)",
            data=docx_file,
            file_name="caca_palavras.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("Por favor, insira ao menos uma palavra.")
