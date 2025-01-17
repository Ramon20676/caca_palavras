import random
import unicodedata
import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt

def normalizar_palavra(palavra):
    # Remove acentos e espaços de uma palavra
    palavra = unicodedata.normalize('NFKD', palavra).encode('ASCII', 'ignore').decode('ASCII')
    return palavra.replace(" ", "")

def verificar_espaco(matriz, palavra, linha, coluna, direcao):
    # Verifica se há espaço disponível para inserir a palavra
    tamanho = len(matriz)
    if direcao == "horizontal":
        if coluna + len(palavra) > tamanho:
            return False
        return all(matriz[linha][coluna + i] in (' ', palavra[i].upper()) for i in range(len(palavra)))
    elif direcao == "vertical":
        if linha + len(palavra) > tamanho:
            return False
        return all(matriz[linha + i][coluna] in (' ', palavra[i].upper()) for i in range(len(palavra)))

def inserir_palavra(matriz, palavra, linha, coluna, direcao):
    # Insere a palavra na matriz na direção especificada
    if direcao == "horizontal":
        for i, letra in enumerate(palavra):
            matriz[linha][coluna + i] = letra.upper()
    elif direcao == "vertical":
        for i, letra in enumerate(palavra):
            matriz[linha + i][coluna] = letra.upper()

def tentar_encaixar_palavras(tamanho, palavras):
    # Criar matriz vazia
    matriz = [[' ' for _ in range(tamanho)] for _ in range(tamanho)]
    coordenadas_palavras = []
    horizontal = True
    for palavra in palavras:
        inserido = False
        for _ in range(100):  # Tentativas para encontrar posição válida
            linha = random.randint(0, tamanho - 1)
            coluna = random.randint(0, tamanho - 1)
            direcao = "horizontal" if horizontal else "vertical"
            if verificar_espaco(matriz, palavra, linha, coluna, direcao):
                inserir_palavra(matriz, palavra, linha, coluna, direcao)
                coordenadas_palavras.append((palavra, linha, coluna, direcao))
                inserido = True
                break
        if not inserido:
            return None, None  # Não foi possível encaixar todas as palavras
        horizontal = not horizontal
    return matriz, coordenadas_palavras

def gerar_caca_palavras(palavras):
    # Determinar tamanho inicial baseado na maior palavra e no número de palavras
    maior_palavra = max(len(palavra) for palavra in palavras)
    tamanho_inicial = max(maior_palavra, len(palavras))

    # Encontrar o menor tamanho possível
    for tamanho in range(tamanho_inicial, 50):  # Limite superior para evitar loops infinitos
        matriz, coordenadas = tentar_encaixar_palavras(tamanho, palavras)
        if matriz is not None:
            break

    # Preencher espaços vazios com letras aleatórias
    for i in range(len(matriz)):
        for j in range(len(matriz[i])):
            if matriz[i][j] == ' ':
                matriz[i][j] = random.choice('ABCDEFGHIJKLMNOPQRSTUVWXYZ')

    return matriz, coordenadas

def salvar_docx(matriz, palavras, coordenadas):
    doc = Document()
    doc.add_heading("Caça-Palavras", level=1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Consolas'
    font.size = Pt(12)

    doc.add_paragraph("Este é o caça-palavras gerado com as palavras fornecidas.")

    doc.add_heading("Versão Original:", level=2)
    for linha in matriz:
        doc.add_paragraph(" ".join(linha))

    # Adicionar a versão destacada em negrito
    doc.add_page_break()
    doc.add_heading("Versão Destacada:", level=2)
    for i in range(len(matriz)):
        p = doc.add_paragraph()
        for j in range(len(matriz[i])):
            run = p.add_run(matriz[i][j] + " ")
            for palavra, linha, coluna, direcao in coordenadas:
                if direcao == "horizontal" and i == linha and coluna <= j < coluna + len(palavra):
                    run.bold = True
                elif direcao == "vertical" and j == coluna and linha <= i < linha + len(palavra):
                    run.bold = True

    # Adicionar menção ao autor no rodapé
    section = doc.sections[-1]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "Gerador criado por [Ramon Costa]"
    footer_paragraph.style = style

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Interface do Streamlit

st.set_page_config(
    page_title="Gerador de Caça-Palavras",
    page_icon=None,  # Substitua por um arquivo de ícone, se desejar
    layout="centered"
)

st.title("Gerador de Caça-Palavras")

st.write("Adicione palavras para criar seu caça-palavras. Quando terminar, clique em 'Gerar'.")

palavras_input = st.text_area("Digite as palavras (uma por linha):")

gerar = st.button("Gerar Caça-Palavras")

if gerar:
    palavras = [normalizar_palavra(p.strip()) for p in palavras_input.split("\n") if p.strip()]
    if palavras:
        matriz, coordenadas = gerar_caca_palavras(palavras)

        st.subheader("Caça-Palavras Gerado:")
        st.code("\n".join([" ".join(linha) for linha in matriz]), language="plaintext")

        docx_file = salvar_docx(matriz, palavras, coordenadas)
        st.download_button(
            label="Baixar Caça-Palavras (.docx)",
            data=docx_file,
            file_name="caca_palavras.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("Por favor, insira ao menos uma palavra.")
